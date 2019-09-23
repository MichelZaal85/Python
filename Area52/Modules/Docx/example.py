# docx example.py

import docx
doc = docx.Document('demo.docx')
print(len(doc.paragraphs))
# 7

print(doc.paragraphs[0].text)
# document title

doc.paragraphs[1].text
# a plain paragraph...

# run is a part of a paragraph / sentence

# A plain paragraph with some bold and some italic
# --------------------------- ---- --------- ------
# run 						   run 	 run  	  run
print(len(doc.paragraphs[1].runs))
# 5

# print(doc.paragraphs[1].runs[0].text)
# print(doc.paragraphs[1].runs[1].text)
# print(doc.paragraphs[1].runs[2].text)
# print(doc.paragraphs[1].runs[3].text)

for i in range(5):
	print(doc.paragraphs[1].runs[i].text)